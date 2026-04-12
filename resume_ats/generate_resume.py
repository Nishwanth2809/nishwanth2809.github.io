"""Generate ATS-friendly resume as a Word document (.docx)"""

from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.0)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(2)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 1.15

ACCENT = RGBColor(0x1E, 0x3A, 0x8A)       # Deep blue
ACCENT_LIGHT = RGBColor(0x3B, 0x82, 0xF6)  # Blue
DARK = RGBColor(0x11, 0x18, 0x27)           # Near black
GRAY = RGBColor(0x4B, 0x55, 0x63)           # Medium gray
LIGHT_GRAY = RGBColor(0x6B, 0x72, 0x80)     # Light gray


def add_colored_line(doc, color_hex="1E3A8A"):
    """Add a colored horizontal line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="{color_hex}"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_section_heading(doc, text):
    """Add a styled section heading."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(text.upper())
    run.font.name = 'Calibri'
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = ACCENT
    run.font.all_caps = True
    add_colored_line(doc)


def add_entry_header(doc, title, date_text=None, link_text=None, link_url=None):
    """Add an entry header with title and optional date."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(1)

    run_title = p.add_run(title)
    run_title.font.name = 'Calibri'
    run_title.font.size = Pt(11)
    run_title.font.bold = True
    run_title.font.color.rgb = DARK

    if link_text:
        run_sep = p.add_run("  ·  ")
        run_sep.font.size = Pt(10)
        run_sep.font.color.rgb = LIGHT_GRAY

        run_link = p.add_run(link_text)
        run_link.font.name = 'Calibri'
        run_link.font.size = Pt(9.5)
        run_link.font.color.rgb = ACCENT_LIGHT
        run_link.font.underline = True

    if date_text:
        # Add date right-aligned using a tab stop
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.4), WD_ALIGN_PARAGRAPH.RIGHT)
        run_tab = p.add_run("\t")
        run_date = p.add_run(date_text)
        run_date.font.name = 'Calibri'
        run_date.font.size = Pt(9.5)
        run_date.font.color.rgb = LIGHT_GRAY
        run_date.font.italic = True


def add_subtitle(doc, text):
    """Add italic subtitle."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = LIGHT_GRAY
    run.font.italic = True


def add_bullet(doc, text):
    """Add a bullet point."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Inches(0.25)

    # Clear default run and add styled text
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(10)
    run.font.color.rgb = GRAY


def add_tech_line(doc, techs):
    """Add a tech stack line with tags."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Inches(0.25)

    run_label = p.add_run("Tech: ")
    run_label.font.name = 'Calibri'
    run_label.font.size = Pt(9)
    run_label.font.bold = True
    run_label.font.color.rgb = ACCENT

    run_items = p.add_run(" · ".join(techs))
    run_items.font.name = 'Calibri'
    run_items.font.size = Pt(9)
    run_items.font.color.rgb = LIGHT_GRAY


# ═══════════════════════════════════════════
# HEADER
# ═══════════════════════════════════════════
name_p = doc.add_paragraph()
name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
name_p.paragraph_format.space_after = Pt(2)
run = name_p.add_run("NISHWANTH YARRA")
run.font.name = 'Calibri'
run.font.size = Pt(24)
run.font.bold = True
run.font.color.rgb = ACCENT

title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(0)
title_p.paragraph_format.space_after = Pt(4)
run = title_p.add_run("AI & Machine Learning Developer  ·  Full-Stack Engineer")
run.font.name = 'Calibri'
run.font.size = Pt(11)
run.font.bold = True
run.font.color.rgb = DARK

contact_p = doc.add_paragraph()
contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
contact_p.paragraph_format.space_before = Pt(0)
contact_p.paragraph_format.space_after = Pt(2)

contact_items = [
    "nishwanthyarra@gmail.com",
    "linkedin.com/in/yarranishwanth",
    "github.com/nishwanth2809",
    "nishwanth2809.github.io",
]

for i, item in enumerate(contact_items):
    run = contact_p.add_run(item)
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.color.rgb = ACCENT_LIGHT if "." in item else GRAY
    if i < len(contact_items) - 1:
        sep = contact_p.add_run("   |   ")
        sep.font.size = Pt(9)
        sep.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)

add_colored_line(doc, "1E3A8A")


# ═══════════════════════════════════════════
# PROFESSIONAL SUMMARY
# ═══════════════════════════════════════════
add_section_heading(doc, "Professional Summary")

summary_p = doc.add_paragraph()
summary_p.paragraph_format.space_after = Pt(4)
run = summary_p.add_run(
    "AI-focused developer with hands-on experience building intelligent systems using "
    "computer vision, NLP, and scalable web technologies. Proven ability to ship end-to-end "
    "products — from training ML models to deploying full-stack applications with real-time "
    "inference. Passionate about turning AI research into deployable tools with clear, "
    "measurable impact."
)
run.font.name = 'Calibri'
run.font.size = Pt(10)
run.font.color.rgb = GRAY


# ═══════════════════════════════════════════
# TECHNICAL SKILLS
# ═══════════════════════════════════════════
add_section_heading(doc, "Technical Skills")

skills_data = [
    ("AI / ML", "Machine Learning, NLP, Computer Vision, Speech Processing (TTS/STT), PyTorch, scikit-learn, Keras, Streamlit"),
    ("Languages", "Python, Java (OOP, DSA), JavaScript, C, SQL, HTML5, CSS3"),
    ("Frontend", "React, Angular.js, React Native, Vite, Tailwind CSS"),
    ("Backend / Cloud", "Node.js, Flask, FastAPI, REST APIs, Firebase, Nginx, Vercel"),
    ("Databases / Data", "MongoDB, SQLite, NumPy, Pandas, Matplotlib, Power BI"),
    ("DevOps / Tools", "Git, GitHub, GitHub Pages, Render, Linux, Postman, VS Code"),
    ("Design", "Figma, Adobe XD, Canva, UI/UX Principles"),
]

table = doc.add_table(rows=len(skills_data), cols=2)
table.alignment = WD_TABLE_ALIGNMENT.LEFT

# Remove borders
for row in table.rows:
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            '  <w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '  <w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            '</w:tcBorders>'
        )
        tcPr.append(tcBorders)

for i, (category, items) in enumerate(skills_data):
    # Category column
    cell_cat = table.cell(i, 0)
    cell_cat.width = Inches(1.3)
    p = cell_cat.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(category)
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.bold = True
    run.font.color.rgb = ACCENT

    # Items column
    cell_items = table.cell(i, 1)
    p = cell_items.paragraphs[0]
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    run = p.add_run(items)
    run.font.name = 'Calibri'
    run.font.size = Pt(9.5)
    run.font.color.rgb = GRAY


# ═══════════════════════════════════════════
# EXPERIENCE
# ═══════════════════════════════════════════
add_section_heading(doc, "Experience")

add_entry_header(doc, "Full Stack Web Development Intern", "Apr 2026 – Oct 2026")
add_subtitle(doc, "ElevanceSkills")
add_bullet(doc, "Developed and deployed real-time web applications using modern frontend and backend technologies in a production environment")
add_bullet(doc, "Collaborated on full-stack projects with deadline-driven deliverables, applying agile development workflows")
add_bullet(doc, "Built responsive user interfaces and integrated RESTful APIs to deliver end-to-end features")

add_entry_header(doc, "Design Member", "2025")
add_subtitle(doc, "SDC — Sreyas Developer's Conference")
add_bullet(doc, "Designed digital content including posters, banners, and event collateral, improving event outreach and engagement")
add_bullet(doc, "Maintained visual consistency across marketing materials by collaborating with cross-functional team members")


# ═══════════════════════════════════════════
# PROJECTS
# ═══════════════════════════════════════════
add_section_heading(doc, "Projects")

add_entry_header(doc, "Deepfake Detection System", link_text="GitHub", link_url="https://github.com/Nishwanth2809/Deep-Fake-Detection")
add_bullet(doc, "Engineered a real-time deepfake detection pipeline using ResNet18 and OpenCV, achieving ~90% classification accuracy")
add_bullet(doc, "Implemented frame-level analysis with live video inference for practical media authenticity verification")
add_bullet(doc, "Deployed an interactive demo using Streamlit for non-technical users to test detection capabilities")
add_tech_line(doc, ["Python", "OpenCV", "ResNet18", "PyTorch", "Streamlit"])

add_entry_header(doc, "AI Medical Report Analyzer", link_text="GitHub", link_url="https://github.com/Nishwanth2809/AI-Medical-Report-Analyzer")
add_bullet(doc, "Built a full-stack application that extracts and analyzes medical reports using OCR and AI-powered NLP workflows")
add_bullet(doc, "Developed a React frontend paired with a Node.js backend to deliver readable AI-generated summaries")
add_bullet(doc, "Simplified complex healthcare report data into actionable patient insights")
add_tech_line(doc, ["React", "Node.js", "OCR", "NLP", "REST APIs"])

add_entry_header(doc, "Mood-Based Music Recommendation (Moodify)", link_text="GitHub", link_url="https://github.com/Nishwanth2809/Moodify-Let-Your-Mood-Choose-the-Music-")
add_bullet(doc, "Developed an NLP-powered mood detection engine that classifies user input and maps it to curated Spotify playlists")
add_bullet(doc, "Integrated Spotify Web API for real-time music recommendations with an interactive experience")
add_tech_line(doc, ["Python", "NLP", "Spotify API", "Streamlit"])

add_entry_header(doc, "AI Attention Monitoring System", link_text="GitHub", link_url="https://github.com/Nishwanth2809/AI-Attention")
add_bullet(doc, "Created a real-time attention tracker using computer vision with gaze detection, blink analysis, and head pose estimation")
add_bullet(doc, "Computed live attention scores and session analytics via Flask backend with MediaPipe and OpenCV")
add_tech_line(doc, ["OpenCV", "MediaPipe", "Flask", "React", "Python"])

add_entry_header(doc, "Voice Chat AI Assistant")
add_bullet(doc, "Built a voice-first AI assistant with real-time speech recognition, NLP, and synthesized responses")
add_bullet(doc, "Developed a fast response loop using FastAPI and speech processing (TTS/STT) for low-latency interaction")
add_tech_line(doc, ["FastAPI", "Python", "TTS/STT"])


# ═══════════════════════════════════════════
# EDUCATION
# ═══════════════════════════════════════════
add_section_heading(doc, "Education")

add_entry_header(doc, "B.Tech in Computer Science (AI & ML Specialization)", "Expected 2027")
add_subtitle(doc, "Sreyas Institute of Engineering and Technology, Hyderabad")


# ═══════════════════════════════════════════
# SAVE
# ═══════════════════════════════════════════
output_path = os.path.join(os.path.dirname(__file__), "Nishwanth_Yarra_Resume.docx")
doc.save(output_path)
print(f"✅ Resume saved: {output_path}")
